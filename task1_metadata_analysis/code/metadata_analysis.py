import os
import sys
import json
import pytesseract
import exifread
from PIL import Image
from langdetect import detect
from deep_translator import GoogleTranslator
from docx import Document

# -----------------------
# Helper: Convert GPS data
# -----------------------
def get_decimal_from_dms(dms, ref):
    degrees = dms[0].num / dms[0].den
    minutes = dms[1].num / dms[1].den
    seconds = dms[2].num / dms[2].den
    decimal = degrees + (minutes / 60.0) + (seconds / 3600.0)
    if ref in ['S', 'W']:
        decimal = -decimal
    return decimal

# -----------------------
# Main analysis
# -----------------------
def analyze_image(image_path, output_path):
    print(f"\n🔍 Analyzing: {image_path}\n")

    exif_data = {}
    gps_data = {}

    try:
        with open(image_path, 'rb') as f:
            tags = exifread.process_file(f, details=False)

        for tag in tags.keys():
            exif_data[tag] = str(tags[tag])

        # Extract GPS if available
        if "GPS GPSLatitude" in tags and "GPS GPSLongitude" in tags:
            lat = get_decimal_from_dms(tags["GPS GPSLatitude"].values, tags["GPS GPSLatitudeRef"].printable)
            lon = get_decimal_from_dms(tags["GPS GPSLongitude"].values, tags["GPS GPSLongitudeRef"].printable)
            gps_data = {"Latitude": lat, "Longitude": lon}
    except Exception as e:
        print(f"[ERROR] Could not extract EXIF: {e}")

    # OCR extraction
    ocr_text = ""
    try:
        ocr_text = pytesseract.image_to_string(Image.open(image_path))
    except Exception as e:
        print(f"[ERROR] OCR extraction failed: {e}")
        ocr_text = ""

    # Language detection
    detected_lang = "Unknown"
    if ocr_text.strip():
        try:
            detected_lang = detect(ocr_text)
        except:
            detected_lang = "Unknown"

    # -----------------------
    # Save output to Word
    # -----------------------
    try:
        doc = Document()
        doc.add_heading("📸 Image Metadata Analysis Report", level=0)

        doc.add_heading("EXIF Metadata", level=1)
        if exif_data:
            for k, v in exif_data.items():
                doc.add_paragraph(f"{k}: {v}")
        else:
            doc.add_paragraph("No EXIF metadata found.")

        doc.add_heading("GPS Data", level=1)
        if gps_data:
            doc.add_paragraph(f"Latitude: {gps_data['Latitude']}")
            doc.add_paragraph(f"Longitude: {gps_data['Longitude']}")
        else:
            doc.add_paragraph("No GPS data found.")

        doc.add_heading("OCR Text Extracted", level=1)
        doc.add_paragraph(ocr_text if ocr_text.strip() else "No text detected.")

        doc.add_heading("Detected Language", level=1)
        doc.add_paragraph(detected_lang)

        doc.save(output_path)
        print(f"✅ Report saved to: {output_path}")

    except Exception as e:
        print(f"[ERROR] Could not save Word file: {e}")

# -----------------------
# Run from CLI
# -----------------------
if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python metadata_analysis.py <image_path>")
        sys.exit(1)

    image_path = sys.argv[1]
    output_dir = os.path.join(os.path.dirname(__file__), "../sample_output")
    os.makedirs(output_dir, exist_ok=True)

    output_file = os.path.join(output_dir, os.path.splitext(os.path.basename(image_path))[0] + "_analysis.docx")
    analyze_image(image_path, output_file)
